import streamlit as st
import os
import re
import torch
import pdfplumber
import pandas as pd
from docx2pdf import convert
from docx.shared import Inches
from docx import Document
from transformers import AutoTokenizer, AutoModelForSequenceClassification
from nltk.tokenize import sent_tokenize
import plotly.express as px
import plotly.io as pio  # For saving Plotly charts as images
import nltk
import tempfile  # For handling temporary files
from io import BytesIO

# Download nltk resource
nltk.download('punkt_tab')

# Streamlit Title and Sidebar Information
st.set_page_config(page_title="SDG Doc Analyzer", page_icon="🌍", layout="wide")

# Define session state initialization
if 'uploaded_file' not in st.session_state:
    st.session_state.uploaded_file = None

if 'df_paragraph_predictions' not in st.session_state:
    st.session_state.df_paragraph_predictions = None

if 'df_sentence_predictions' not in st.session_state:
    st.session_state.df_sentence_predictions = None

# Function to clear session state
def clear_analysis():
    st.session_state.uploaded_file = None
    st.session_state.df_paragraph_predictions = None
    st.session_state.df_sentence_predictions = None
    st.rerun()  # This will trigger a rerun of the app

# Sidebar to clear all records and start new analysis
if st.sidebar.button("Clear and Start New Analysis"):
    clear_analysis()

# Model checkpoint
checkpoint = "sadickam/sdgBERT"

# Define keywords/phrases for sections to exclude
EXCLUDED_SECTION_KEYWORDS = [
    r'table\s+of\s+contents', r'list\s+of\s+figures', r'list\s+of\s+tables', r'abbreviations',
    r'references', r'bibliography', r'glossary', r'appendix', r'contents', r'content',
    r'list\s+of\s+definitions', r'definitions', r'tables', r'figures'
]

# Define patterns for identifying section start and headings
HEADING_PATTERNS = [
    r'^[A-Z][A-Za-z\s]{0,100}$',  # Title-like headings (first letter capitalized, up to 100 chars)
    r'^\d+(\.\d+)*\s+[A-Z][A-Za-z\s]{0,100}$',  # Numbered headings (e.g., 1. Introduction)
    r'^[IVXLC]+\.\s+[A-Z][A-Za-z\s]+$',  # Roman numeral headings (e.g., IV. Analysis)
    r'^[A-Za-z]\.\s+[A-Z][A-Za-z\s]+$',  # Letter outline (e.g., A. Definitions)
    r'.+(\.+|\s)\d+$',  # Text with dot leaders or page numbers (common in ToC)
]

# Define a threshold for short text (common in headings)
HEADING_MAX_WORDS = 10

# Function to detect if text is part of excluded sections
def is_excluded_section(text):
    text_lower = text.lower()
    for keyword in EXCLUDED_SECTION_KEYWORDS:
        if re.search(keyword, text_lower):
            return True

    # Check for heading patterns (short, possibly capitalized text)
    if len(text.split()) <= HEADING_MAX_WORDS:
        for pattern in HEADING_PATTERNS:
            if re.match(pattern, text.strip()):
                return True
    return False

# Function to clean extracted text
def clean_text(text):
    text = text.strip()
    text = re.sub(r'\s+', ' ', text)
    return text

# Function to extract paragraphs and tables from PDF while excluding unwanted sections
def extract_pdf_content(pdf_file):
    document_name = pdf_file.name
    df = pd.DataFrame(columns=["document_name", "page_number", "paragraph_number", "row_number", "content_type", "text"])

    with pdfplumber.open(pdf_file) as pdf:
        paragraph_number = 0
        row_number = 0
        exclude_mode = False  # Flag to indicate if we are in an excluded section

        for page_num, page in enumerate(pdf.pages, start=1):
            # Extract paragraphs
            text = clean_text(page.extract_text())
            paragraphs = text.split('\n')

            for paragraph in paragraphs:
                cleaned_paragraph = clean_text(paragraph)

                # Check if the paragraph is part of an excluded section
                if is_excluded_section(cleaned_paragraph):
                    exclude_mode = True

                # If the excluded section is detected, skip paragraphs until the next section is found
                if exclude_mode:
                    if len(cleaned_paragraph.split()) > 5 and not is_excluded_section(cleaned_paragraph):
                        exclude_mode = False  # End of excluded section
                    else:
                        continue  # Skip current paragraph

                if cleaned_paragraph and len(cleaned_paragraph.split()) > 2:  # Valid paragraph check
                    paragraph_number += 1
                    df = pd.concat([df, pd.DataFrame([{
                        "document_name": document_name,
                        "page_number": page_num,
                        "paragraph_number": paragraph_number,
                        "row_number": None,
                        "content_type": "paragraph",
                        "text": cleaned_paragraph,
                    }])], ignore_index=True)

            # Extract tables
            tables = page.extract_tables()
            for table in tables:
                for row in table:
                    row_number += 1
                    row_text = " | ".join([str(cell).strip() for cell in row if cell])  # Join all cells in the row
                    if row_text:  # Ensure it's not empty
                        df = pd.concat([df, pd.DataFrame([{
                            "document_name": document_name,
                            "page_number": page_num,
                            "paragraph_number": None,
                            "row_number": row_number,
                            "content_type": "table_row",
                            "text": row_text,
                        }])], ignore_index=True)

    return df

# Function to convert DOCX to PDF and return the PDF path
def convert_docx_to_pdf(docx_file):
    # Save the uploaded file to a temporary location
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_docx:
        tmp_docx.write(docx_file.getbuffer())
        tmp_docx_path = tmp_docx.name

    # Convert DOCX to PDF
    tmp_pdf_path = tmp_docx_path.replace('.docx', '.pdf')
    convert(tmp_docx_path, tmp_pdf_path)

    return tmp_pdf_path

# Function to extract content based on file type
def extract_content(file):
    file_extension = os.path.splitext(file.name)[1].lower()

    if file_extension == '.pdf':
        return extract_pdf_content(file)
    elif file_extension == '.docx':
        # Convert DOCX to PDF
        pdf_path = convert_docx_to_pdf(file)
        with open(pdf_path, 'rb') as pdf_file:
            return extract_pdf_content(pdf_file)
    else:
        return pd.DataFrame()

# Preprocessing function for text
def prep_text(text):
    clean_sents = []
    sent_tokens = sent_tokenize(str(text))
    for sent_token in sent_tokens:
        word_tokens = [str(word_token).strip().lower() for word_token in sent_token.split()]
        clean_sents.append(' '.join(word_tokens))
    joined = ' '.join(clean_sents).strip()
    return re.sub(r'`|"', "", joined)

# Load the tokenizer and model with GPU support
@st.cache_resource
def load_model_and_tokenizer():
    model = AutoModelForSequenceClassification.from_pretrained(checkpoint).to(device)
    tokenizer = AutoTokenizer.from_pretrained(checkpoint)
    return model, tokenizer

# Define device (ensure usage of GPU if available in Hugging Face Spaces)
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

# SDG labels
label_list = [
    'SDG1_No Poverty', 'SDG2_Zero Hunger', 'SDG3_Good Health and Well-being', 'SDG4_Quality Education',
    'SDG5_Gender Equality', 'SDG6_Clean Water and Sanitation', 'SDG7_Affordable and Clean Energy',
    'SDG8_Decent Work and Economic Growth', 'SDG9_Industry, Innovation and Infrastructure',
    'SDG10_Reduced Inequality', 'SDG11_Sustainable Cities and Communities',
    'SDG12_Responsible Consumption and Production', 'SDG13_Climate Action',
    'SDG14_Life Below Water', 'SDG15_Life on Land', 'SDG16_Peace, Justice and Strong Institutions'
]

# Function to predict SDGs for a given text input
def predict_sdg_labels(text, model, tokenizer):
    tokenized_text = tokenizer(text, return_tensors="pt", truncation=True, max_length=512).to(device)
    model.eval()
    with torch.no_grad():
        text_logits = model(**tokenized_text).logits
    predictions = torch.softmax(text_logits, dim=1).tolist()[0]
    return predictions

# Paragraph-level predictions
def predict_paragraphs(df_combined_paragraphs):
    model, tokenizer = load_model_and_tokenizer()

    all_predicted_labels = [[] for _ in range(16)]
    all_prediction_scores = [[] for _ in range(16)]

    for text_input in df_combined_paragraphs['text']:
        joined_clean_sents = prep_text(text_input)
        predictions = predict_sdg_labels(joined_clean_sents, model, tokenizer)
        sorted_preds = sorted(zip(label_list, predictions), key=lambda x: x[1], reverse=True)

        for i, (label, score) in enumerate(sorted_preds):
            all_predicted_labels[i].append(label)
            all_prediction_scores[i].append(score)

    for i in range(16):
        df_combined_paragraphs[f'pred{i + 1}'] = all_predicted_labels[i]
        df_combined_paragraphs[f'score{i + 1}'] = all_prediction_scores[i]

    return df_combined_paragraphs

# Sentence-level predictions
def predict_sentences(df_combined_paragraphs):
    model, tokenizer = load_model_and_tokenizer()
    df_combined_sentences = pd.DataFrame()

    for _, row in df_combined_paragraphs.iterrows():
        sentences = sent_tokenize(row["text"])
        for sentence in sentences:
            clean_sentence = prep_text(sentence)
            predictions = predict_sdg_labels(clean_sentence, model, tokenizer)
            sorted_predictions = sorted(zip(label_list, predictions), key=lambda x: x[1], reverse=True)
            pred_labels, pred_scores = zip(*sorted_predictions)
            sentence_row = {
                "document_name": row["document_name"],
                "page_number": row.get("page_number", None),
                "paragraph_number": row["paragraph_number"],
                "text": sentence,
                **{f"pred{i}": pred_labels[i-1] for i in range(1, 17)},
                **{f"score{i}": round(pred_scores[i-1], 3) for i in range(1, 17)}
            }
            df_combined_sentences = pd.concat([df_combined_sentences, pd.DataFrame([sentence_row])], ignore_index=True)

    return df_combined_sentences

# Updated Plotting function to create First, Second, and Third Dominant SDGs with percentages and colorful bars
def plot_sdg_dominant(df, title, pred_column):
    df_filtered = df[df[f'score{pred_column[-1]}'] > 0]
    labels = df_filtered[pred_column].value_counts()

    # Calculate percentages instead of raw counts
    total_count = labels.sum()
    percentages = (labels / total_count) * 100

    # Create a bar plot with percentages and colorful bars
    fig = px.bar(
        percentages.rename_axis('SDG Label').reset_index(name='Percentage'), 
        y='SDG Label', 
        x='Percentage', 
        orientation='h', 
        title=title,
        color='SDG Label',  # Colorful bars based on the SDG label
        color_discrete_sequence=px.colors.qualitative.Plotly  # Set a colorful theme
    )

    # Adjusting layout to ensure labels aren't cut off and display percentages
    fig.update_layout(
        yaxis=dict(
            automargin=True,  # Ensures enough space for long labels
            tickmode='array',  # Fix tick mode
            title=None  # Turn off y-axis title
        ),
        margin=dict(l=150, r=20, t=60, b=50),  # Adjusting margins
        height=500,  # Set a larger height to accommodate long labels
        showlegend=False  # Turn off the legend
    )
    
    # Adjust the hover information to show percentages more clearly
    fig.update_traces(
        hovertemplate='%{y}: %{x:.2f}%',  # Show percentage in hover label
        texttemplate='%{x:.2f}%',         # Show percentage inside the bars
        textposition='auto'               # Automatically position the text
    )
    
    return fig

# Function to create CSV data for download
def create_csv(df, file_name):
    return df.to_csv(index=False).encode('utf-8')

# Save and return the path to the plot images
def save_plot_as_image(fig, filename):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_file:
        fig.write_image(temp_file.name)
        return temp_file.name

# Function to create DOCX report and include plots based on the type of analysis
def generate_docx_report(df_paragraphs, df_sentences, paragraph_plots, sentence_plots, analysis_type):
    doc = Document()
    doc.add_heading("SDG Analysis Report", 0)

    if analysis_type == "Paragraph-Level":
        # Add paragraph-level analysis
        doc.add_heading("Paragraph-Level SDG Analysis", level=1)
        if not df_paragraphs.empty:
            doc.add_heading("Interpreting Bar Graphs", level=2)
            doc.add_paragraph(
                "Each paragraph is processed by an AI model trained to predict which of the first 16 Sustainable Development Goals (SDGs) "
                "is most relevant to the text. The model analyzes the content and returns a prediction for each SDG, providing scores that "
                "represent the likelihood that the text is related to a particular SDG. The predictions are based on probabilities, and "
                "the model outputs predictions for the first 16 SDGs. This paragraph-level analysis provides high-level insight into SDG alignment, "
                "focusing on the top three SDG predictions for each paragraph with a probability score greater than zero."
            )

            doc.add_paragraph(
                "1. First Dominant SDGs Bar Graph: This graph displays the primary SDG that the AI model associates with each text. "
                "The bars represent the percentage of paragraphs that are most strongly aligned with each SDG. This is the strongest "
                "indication of relevance for each text, offering insight into the dominant sustainable development theme within the text."
            )

            doc.add_paragraph(
                "2. Second Dominant SDGs Bar Graph: This graph shows the second most relevant SDG for each text, where the model predicts "
                "that although this SDG is not the primary focus, the text still has significant relevance to this goal."
            )

            doc.add_paragraph(
                "3. Third Dominant SDGs Bar Graph: Similar to the previous graphs, this graph represents the third most relevant SDG for each "
                "text. It provides further insight into the text's alignment with multiple SDGs, offering a broader understanding of the "
                "content's focus areas."
            )
            for doc_name in df_paragraphs['document_name'].unique():
                doc.add_heading(f"Document: {doc_name}", level=2)
                doc.add_heading("First Dominant SDGs", level=3)

                # Add paragraph-level plots to the DOCX
                for plot in paragraph_plots:
                    doc.add_picture(plot, width=Inches(6))  # Add plot image to DOCX
        else:
            doc.add_paragraph("No paragraph-level analysis available.")

    elif analysis_type == "Sentence-Level":
        # Add sentence-level analysis
        doc.add_heading("Sentence-Level SDG Analysis", level=1)
        if not df_sentences.empty:
            doc.add_heading("Interpreting Bar Graphs", level=1)
            doc.add_paragraph(
                "Each sentence is processed by an AI model trained to predict which of the first 16 Sustainable Development Goals (SDGs) "
                "is most relevant to the text. The model analyzes the content and returns a prediction for each SDG, providing scores that "
                "represent the likelihood that the text is related to a particular SDG. The predictions are based on probabilities, and "
                "the model outputs predictions for the first 16 SDGs. This sentence-level analysis provides deeper insight into SDG alignment, "
                "focusing on the top three SDG predictions for each sentence with a probability score greater than zero."
            )

            doc.add_paragraph(
                "1. First Dominant SDGs Bar Graph: This graph displays the primary SDG that the AI model associates with each sentence. "
                "The bars represent the percentage of sentences that are most strongly aligned with each SDG. This is the strongest "
                "indication of relevance for each sentence, offering insight into the dominant sustainable development theme within the text."
            )

            doc.add_paragraph(
                "2. Second Dominant SDGs Bar Graph: This graph shows the second most relevant SDG for each sentence, where the model predicts "
                "that although this SDG is not the primary focus, the text still has significant relevance to this goal."
            )

            doc.add_paragraph(
                "3. Third Dominant SDGs Bar Graph: Similar to the previous graphs, this graph represents the third most relevant SDG for each "
                "sentence. It provides further insight into the text's alignment with multiple SDGs, offering a broader understanding of the "
                "content's focus areas."
            )
            
            for doc_name in df_sentences['document_name'].unique():
                doc.add_heading(f"Document: {doc_name}", level=2)
                doc.add_heading("First Dominant SDGs", level=3)

                # Add sentence-level plots to the DOCX
                for plot in sentence_plots:
                    doc.add_picture(plot, width=Inches(6))  # Add plot image to DOCX
        else:
            doc.add_paragraph("No sentence-level analysis available.")

    # Save DOCX to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit Interface

st.sidebar.title("SDG Document Analysis")
st.sidebar.write("Upload a PDF or DOCX file to analyse its alignment with the UN Sustainable Development Goals (SDGs).")

uploaded_file = st.sidebar.file_uploader("Upload PDF/DOCX File", type=["pdf", "docx"])

# Improved Layout
st.title("SDG Document Analysis App 🌍")
st.markdown("""
    This app allows you to upload **PDF** and **DOCX** files to analyze their alignment with the 
    [United Nations Sustainable Development Goals](https://sdgs.un.org/goals). You can generate reports for paragraph-level 
    and sentence-level predictions, visualize the most relevant SDGs, and download CSV and DOCX reports. 
    
    This app breaks a document into paragraphs by detecting new lines. Hence, a sentence may be identified as a paragraph
    if a new line is introduced after it. Table text is extracted in rows, and the content for each row is concatenated.
    The analysis and the CSV reports exclude the table of contents, list of abbreviations, and reference lists.
""")

if uploaded_file is not None:
    st.subheader("File Summary")
    st.write(f"**File name:** {uploaded_file.name}")
    st.write(f"**File size:** {uploaded_file.size / 1024:.2f} KB")

    with st.spinner('Processing file...'):
        df_paragraphs = extract_content(uploaded_file)

        # **Optimization:**
        # Compute df_paragraph_predictions only once, outside the if-elif block.
        if not df_paragraphs.empty:
            df_paragraph_predictions = predict_paragraphs(df_paragraphs)
        else:
            st.error("No valid content found in the uploaded file.")
            st.stop()

        # Add Option to Select Type of Analysis
        analysis_type = st.radio("Choose Analysis Type", ["Paragraph-Level", "Sentence-Level"])

        # Layout Columns for Plots
        col1, col2, col3 = st.columns(3)

        paragraph_plots = []
        sentence_plots = []

        if analysis_type == "Paragraph-Level":
            st.subheader("Paragraph-Level SDG Predictions")

            with col1:
                first_sdg_paragraph = plot_sdg_dominant(df_paragraph_predictions, "First Dominant SDGs", 'pred1')
                st.plotly_chart(first_sdg_paragraph, use_container_width=True)
                st.write("""
                This graph displays the primary SDGs that the AI model associates with each paragraph. The bars represent the percentage of 
                paragraphs most strongly aligned with each SDG. This is the strongest indication of relevance for each text, 
                offering insight into the dominant sustainable development theme within the text.
                """)
                paragraph_plots.append(save_plot_as_image(first_sdg_paragraph, "paragraph_first_sdg.png"))

            with col2:
                second_sdg_paragraph = plot_sdg_dominant(df_paragraph_predictions, "Second Dominant SDGs", 'pred2')
                st.plotly_chart(second_sdg_paragraph, use_container_width=True)
                st.write("""
                This graph shows the second most relevant SDG for each paragraph, where the model predicts that although 
                this SDG is not the primary focus, but the text is still relevant to this goal.
                """)
                paragraph_plots.append(save_plot_as_image(second_sdg_paragraph, "paragraph_second_sdg.png"))

            with col3:
                third_sdg_paragraph = plot_sdg_dominant(df_paragraph_predictions, "Third Dominant SDGs", 'pred3')
                st.plotly_chart(third_sdg_paragraph, use_container_width=True)
                st.write("""This graph represents the third most relevant SDG for each paragraph. It provides further insight into the text's alignment
                with multiple SDGs, offering a broader understanding of the content's focus areas.""")
                paragraph_plots.append(save_plot_as_image(third_sdg_paragraph, "paragraph_third_sdg.png"))

            # Provide paragraph-level CSV download
            st.download_button("Download Paragraph Predictions CSV", data=create_csv(df_paragraph_predictions, 'paragraph_predictions.csv'),
                               file_name='paragraph_predictions.csv')

        elif analysis_type == "Sentence-Level":
            st.subheader("Sentence-Level SDG Predictions")

            # No need to recompute df_paragraph_predictions; it's already done.
            df_sentence_predictions = predict_sentences(df_paragraph_predictions)

            # Display sentence-level predictions with three bar plots
            with col1:
                first_sdg_sentence = plot_sdg_dominant(df_sentence_predictions, "Sentence: First Dominant SDGs", 'pred1')
                st.plotly_chart(first_sdg_sentence, use_container_width=True)
                st.write("""
                This graph displays the primary SDGs that the AI model associates with each sentence. The bars represent the percentage of 
                sentences most strongly aligned with each SDG. This is the strongest indication of relevance for each text, 
                offering insight into the dominant sustainable development theme within the text.
                """)
                sentence_plots.append(save_plot_as_image(first_sdg_sentence, "sentence_first_sdg.png"))

            with col2:
                second_sdg_sentence = plot_sdg_dominant(df_sentence_predictions, "Sentence: Second Dominant SDGs", 'pred2')
                st.plotly_chart(second_sdg_sentence, use_container_width=True)
                st.write("""
                This graph shows the second most relevant SDG for each sentence, where the model predicts that although 
                this SDG is not the primary focus, but the text is still relevant to this goal.
                """)
                sentence_plots.append(save_plot_as_image(second_sdg_sentence, "sentence_second_sdg.png"))

            with col3:
                third_sdg_sentence = plot_sdg_dominant(df_sentence_predictions, "Sentence: Third Dominant SDGs", 'pred3')
                st.plotly_chart(third_sdg_sentence, use_container_width=True)
                st.write("""This graph represents the third most relevant SDG for each sentence. It provides further insight into the text's alignment
                with multiple SDGs, offering a broader understanding of the content's focus areas.""")
                sentence_plots.append(save_plot_as_image(third_sdg_sentence, "sentence_third_sdg.png"))

            # Provide sentence-level CSV download
            st.download_button("Download Sentence Predictions CSV", data=create_csv(df_sentence_predictions, 'sentence_predictions.csv'),
                               file_name='sentence_predictions.csv')

        # Generate and download DOCX report (only available if both analyses are run)
        st.markdown("---")
        st.subheader("Download Reports")

        # Generate and download DOCX report based on the selected analysis type
        if analysis_type == "Paragraph-Level" or analysis_type == "Sentence-Level":
            buffer = generate_docx_report(
                df_paragraph_predictions, locals().get('df_sentence_predictions', pd.DataFrame()), paragraph_plots, sentence_plots, analysis_type
            )
            st.download_button("Download DOCX Report", data=buffer, file_name="sdg_analysis_report.docx")

else:
    st.info("Please upload a PDF or DOCX file to begin the analysis.")
