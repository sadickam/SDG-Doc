import streamlit as st
import os
import re
import torch
import pdfplumber
import pandas as pd
from docx import Document
from docx.shared import Inches
from transformers import AutoTokenizer, AutoModelForSequenceClassification
from nltk.tokenize import sent_tokenize
import plotly.express as px
import plotly.io as pio  # For saving Plotly charts as images
import nltk
from io import BytesIO
import tempfile  # For temporary files

# Download nltk resource
nltk.download('punkt_tab')

# Model checkpoint for SDG BERT
checkpoint = "sadickam/sdgBERT"

# Function to clean extracted text
def clean_text(text):
    text = text.strip()
    text = re.sub(r'\s+', ' ', text)
    return text

# Function to extract paragraphs from PDF
def extract_pdf_content(pdf_file):
    document_name = pdf_file.name
    df = pd.DataFrame(columns=["document_name", "page_number", "paragraph_number", "text"])

    with pdfplumber.open(pdf_file) as pdf:
        paragraph_number = 0
        for page_num, page in enumerate(pdf.pages, start=1):
            text = clean_text(page.extract_text())
            paragraphs = text.split('\n')

            for paragraph in paragraphs:
                cleaned_paragraph = clean_text(paragraph)
                if cleaned_paragraph and len(cleaned_paragraph.split()) > 2:  # Valid paragraph check
                    paragraph_number += 1
                    df = pd.concat([df, pd.DataFrame([{
                        "document_name": document_name,  # Ensure document name is added
                        "page_number": page_num,
                        "paragraph_number": paragraph_number,
                        "text": cleaned_paragraph,
                    }])], ignore_index=True)

    return df

# Function to extract paragraphs from DOCX
def extract_docx_content(docx_file):
    document_name = docx_file.name
    doc = Document(docx_file)
    df = pd.DataFrame(columns=["document_name", "paragraph_number", "text"])

    paragraph_number = 0
    for para in doc.paragraphs:
        cleaned_paragraph = clean_text(para.text)
        if cleaned_paragraph and len(cleaned_paragraph.split()) > 2:  # Valid paragraph check
            paragraph_number += 1
            df = pd.concat([df, pd.DataFrame([{
                "document_name": document_name,  # Ensure document name is added
                "paragraph_number": paragraph_number,
                "text": cleaned_paragraph
            }])], ignore_index=True)

    return df

# Function to extract content based on file type
def extract_content(file):
    file_extension = os.path.splitext(file.name)[1].lower()

    if file_extension == '.pdf':
        return extract_pdf_content(file)
    elif file_extension == '.docx':
        return extract_docx_content(file)
    else:
        return pd.DataFrame()

# Preprocessing function for text
def prep_text(text):
    clean_sents = []
    sent_tokens = sent_tokenize(str(text))
    for sent_token in sent_tokens:
        word_tokens = [str(word_token).strip().lower() for word_token in sent_token.split()]
        clean_sents.append(' '.join(word_tokens))
    joined = ' '.join(clean_sents).strip()
    return re.sub(r'`|"', "", joined)

# Load the tokenizer and model with GPU support
@st.cache_resource
def load_model_and_tokenizer():
    model = AutoModelForSequenceClassification.from_pretrained(checkpoint).to(device)
    tokenizer = AutoTokenizer.from_pretrained(checkpoint)
    return model, tokenizer

# Define device (ensure usage of GPU if available in Hugging Face Spaces)
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")

# SDG labels
label_list = [
    'SDG1_No Poverty', 'SDG2_Zero Hunger', 'SDG3_Good Health and Well-being', 'SDG4_Quality Education',
    'SDG5_Gender Equality', 'SDG6_Clean Water and Sanitation', 'SDG7_Affordable and Clean Energy',
    'SDG8_Decent Work and Economic Growth', 'SDG9_Industry, Innovation and Infrastructure',
    'SDG10_Reduced Inequality', 'SDG11_Sustainable Cities and Communities',
    'SDG12_Responsible Consumption and Production', 'SDG13_Climate Action',
    'SDG14_Life Below Water', 'SDG15_Life on Land', 'SDG16_Peace, Justice and Strong Institutions'
]

# Function to predict SDGs for a given text input
def predict_sdg_labels(text, model, tokenizer):
    tokenized_text = tokenizer(text, return_tensors="pt", truncation=True, max_length=512).to(device)
    model.eval()
    with torch.no_grad():
        text_logits = model(**tokenized_text).logits
    predictions = torch.softmax(text_logits, dim=1).tolist()[0]
    return predictions

# Paragraph-level predictions
def predict_paragraphs(df_combined_paragraphs):
    model, tokenizer = load_model_and_tokenizer()

    all_predicted_labels = [[] for _ in range(16)]
    all_prediction_scores = [[] for _ in range(16)]

    for text_input in df_combined_paragraphs['text']:
        joined_clean_sents = prep_text(text_input)
        predictions = predict_sdg_labels(joined_clean_sents, model, tokenizer)
        sorted_preds = sorted(zip(label_list, predictions), key=lambda x: x[1], reverse=True)

        for i, (label, score) in enumerate(sorted_preds):
            all_predicted_labels[i].append(label)
            all_prediction_scores[i].append(score)

    for i in range(16):
        df_combined_paragraphs[f'pred{i + 1}'] = all_predicted_labels[i]
        df_combined_paragraphs[f'score{i + 1}'] = all_prediction_scores[i]

    return df_combined_paragraphs

# Sentence-level predictions
def predict_sentences(df_combined_paragraphs):
    model, tokenizer = load_model_and_tokenizer()
    df_combined_sentences = pd.DataFrame()

    for _, row in df_combined_paragraphs.iterrows():
        sentences = sent_tokenize(row["text"])
        for sentence in sentences:
            clean_sentence = prep_text(sentence)
            predictions = predict_sdg_labels(clean_sentence, model, tokenizer)
            sorted_predictions = sorted(zip(label_list, predictions), key=lambda x: x[1], reverse=True)
            pred_labels, pred_scores = zip(*sorted_predictions)
            sentence_row = {
                "document_name": row["document_name"],  # Ensure document name is propagated
                "page_number": row.get("page_number", None),
                "paragraph_number": row["paragraph_number"],
                "text": sentence,
                **{f"pred{i}": pred_labels[i-1] for i in range(1, 17)},
                **{f"score{i}": round(pred_scores[i-1], 3) for i in range(1, 17)}
            }
            df_combined_sentences = pd.concat([df_combined_sentences, pd.DataFrame([sentence_row])], ignore_index=True)

    return df_combined_sentences

# Plotting function to create First, Second, and Third Dominant SDGs
def plot_sdg_dominant(df, title, pred_column):
    df_filtered = df[df[f'score{pred_column[-1]}'] > 0]
    labels = df_filtered[pred_column].value_counts()
    fig = px.bar(labels.rename_axis('SDG Label').reset_index(name='Count'), 
                 y='SDG Label', x='Count', orientation='h', title=title)
    return fig

# Function to create CSV data for download
def create_csv(df, file_name):
    return df.to_csv(index=False).encode('utf-8')

# Save and return the path to the plot images
def save_plot_as_image(fig, filename):
    with tempfile.NamedTemporaryFile(delete=False, suffix=".png") as temp_file:
        fig.write_image(temp_file.name)
        return temp_file.name

# Function to create DOCX report and include plots
def generate_docx_report(df_paragraphs, df_sentences, paragraph_plots, sentence_plots):
    doc = Document()
    doc.add_heading("SDG Analysis Report", 0)

    # Add paragraph-level analysis
    doc.add_heading("Paragraph-Level SDG Analysis", level=1)
    if not df_paragraphs.empty:
        for doc_name in df_paragraphs['document_name'].unique():
            doc.add_heading(f"Document: {doc_name}", level=2)
            doc.add_heading("First Dominant SDGs", level=3)

            # Add paragraph-level plots to the DOCX
            for plot in paragraph_plots:
                doc.add_picture(plot, width=Inches(6))  # Add plot image to DOCX
    else:
        doc.add_paragraph("No paragraph-level analysis available.")

    # Add sentence-level analysis
    doc.add_heading("Sentence-Level SDG Analysis", level=1)
    if not df_sentences.empty:
        for doc_name in df_sentences['document_name'].unique():
            doc.add_heading(f"Document: {doc_name}", level=2)
            doc.add_heading("First Dominant SDGs", level=3)

            # Add sentence-level plots to the DOCX
            for plot in sentence_plots:
                doc.add_picture(plot, width=Inches(6))  # Add plot image to DOCX
    else:
        doc.add_paragraph("No sentence-level analysis available.")

    # Save DOCX to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Streamlit Interface

# Title and Sidebar Information
st.set_page_config(page_title="SDG Predictor", page_icon="🌍", layout="wide")

st.sidebar.title("SDG File Analysis")
st.sidebar.write("Upload a PDF or DOCX file to predict its alignment with the UN Sustainable Development Goals (SDGs).")

uploaded_file = st.sidebar.file_uploader("Upload PDF/DOCX File", type=["pdf", "docx"])

# Improved Layout
st.title("Sustainable Development Goal (SDG) Prediction App 🌍")
st.markdown("""
    This app allows you to upload **PDF** and **DOCX** files to analyze their alignment with the 
    [United Nations Sustainable Development Goals](https://sdgs.un.org/goals). You can generate reports for paragraph-level 
    and sentence-level predictions, visualize the most relevant SDGs, and download CSV and DOCX reports.
""")

if uploaded_file is not None:
    st.subheader("File Summary")
    st.write(f"**File name:** {uploaded_file.name}")
    st.write(f"**File size:** {uploaded_file.size / 1024:.2f} KB")

    with st.spinner('Processing file...'):
        df_paragraphs = extract_content(uploaded_file)
        df_paragraph_predictions = predict_paragraphs(df_paragraphs)

        # Add Option to Select Type of Analysis
        analysis_type = st.radio("Choose Analysis Type", ["Paragraph-Level", "Sentence-Level"])

        # Layout Columns for Plots
        col1, col2, col3 = st.columns(3)

        paragraph_plots = []
        sentence_plots = []

        if analysis_type == "Paragraph-Level":
            st.subheader("Paragraph-Level SDG Predictions")
            with col1:
                first_sdg_paragraph = plot_sdg_dominant(df_paragraph_predictions, "First Dominant SDGs", 'pred1')
                st.plotly_chart(first_sdg_paragraph, use_container_width=True)
                paragraph_plots.append(save_plot_as_image(first_sdg_paragraph, "paragraph_first_sdg.png"))

            with col2:
                second_sdg_paragraph = plot_sdg_dominant(df_paragraph_predictions, "Second Dominant SDGs", 'pred2')
                st.plotly_chart(second_sdg_paragraph, use_container_width=True)
                paragraph_plots.append(save_plot_as_image(second_sdg_paragraph, "paragraph_second_sdg.png"))

            with col3:
                third_sdg_paragraph = plot_sdg_dominant(df_paragraph_predictions, "Third Dominant SDGs", 'pred3')
                st.plotly_chart(third_sdg_paragraph, use_container_width=True)
                paragraph_plots.append(save_plot_as_image(third_sdg_paragraph, "paragraph_third_sdg.png"))

            # Provide paragraph-level CSV download
            st.download_button("Download Paragraph Predictions CSV", data=create_csv(df_paragraph_predictions, 'paragraph_predictions.csv'),
                               file_name='paragraph_predictions.csv')

        elif analysis_type == "Sentence-Level":
            st.subheader("Sentence-Level SDG Predictions")
            df_sentence_predictions = predict_sentences(df_paragraph_predictions)

            # Display sentence-level predictions with three bar plots
            with col1:
                first_sdg_sentence = plot_sdg_dominant(df_sentence_predictions, "Sentence: First Dominant SDGs", 'pred1')
                st.plotly_chart(first_sdg_sentence, use_container_width=True)
                sentence_plots.append(save_plot_as_image(first_sdg_sentence, "sentence_first_sdg.png"))

            with col2:
                second_sdg_sentence = plot_sdg_dominant(df_sentence_predictions, "Sentence: Second Dominant SDGs", 'pred2')
                st.plotly_chart(second_sdg_sentence, use_container_width=True)
                sentence_plots.append(save_plot_as_image(second_sdg_sentence, "sentence_second_sdg.png"))

            with col3:
                third_sdg_sentence = plot_sdg_dominant(df_sentence_predictions, "Sentence: Third Dominant SDGs", 'pred3')
                st.plotly_chart(third_sdg_sentence, use_container_width=True)
                sentence_plots.append(save_plot_as_image(third_sdg_sentence, "sentence_third_sdg.png"))

            # Provide sentence-level CSV download
            st.download_button("Download Sentence Predictions CSV", data=create_csv(df_sentence_predictions, 'sentence_predictions.csv'),
                               file_name='sentence_predictions.csv')

        # Generate and download DOCX report (only available if both analyses are run)
        st.markdown("---")
        st.subheader("Download Reports")

        # Check if both analyses were performed
        if analysis_type == "Paragraph-Level" or 'df_sentence_predictions' in locals():
            buffer = generate_docx_report(df_paragraph_predictions, locals().get('df_sentence_predictions', pd.DataFrame()), paragraph_plots, sentence_plots)
            st.download_button("Download DOCX Report", data=buffer, file_name="sdg_analysis_report.docx")

else:
    st.info("Please upload a PDF or DOCX file to begin the analysis.")
